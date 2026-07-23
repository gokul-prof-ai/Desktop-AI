"""
DesktopAI
DOCX Reader

Extracts text content from Word (.docx) files, including both
regular paragraphs and any tables in the document.
"""

import zipfile
from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

from core.logger import get_logger

logger = get_logger("documents")


def read_docx_text(path: Path) -> str | None:
    """
    Extract all text from a DOCX file, including tables.

    Args:
        path (Path):
            The DOCX file to read.

    Returns:
        str | None:
            The extracted text, or None if the file could not be
            read (missing, corrupted, or not a valid DOCX file).
            A None result is expected behavior for a bad file,
            not a crash.
    """

    if not path.exists():
        logger.warning("DOCX not found: %s", path)
        return None

    try:
        document = docx.Document(path)
    except (PackageNotFoundError, KeyError, zipfile.BadZipFile) as error:
        # PackageNotFoundError: file is missing/empty/not a zip at all.
        # KeyError: file is a valid zip, but not a valid DOCX structure
        # (missing the internal [Content_Types].xml part).
        # BadZipFile: file isn't a valid zip archive at all (since a
        # DOCX file is itself a zip archive internally).
        logger.warning("Could not open DOCX %s: %s", path, error)
        return None

    text_parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    text = "\n".join(part for part in text_parts if part.strip())

    logger.info("Read %d character(s) from %s", len(text), path)

    return text