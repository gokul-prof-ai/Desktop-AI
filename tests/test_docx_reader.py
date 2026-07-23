"""
DesktopAI
Tests for the DOCX Reader module.
"""

import docx

from documents.docx_reader import read_docx_text


def make_test_docx(path, paragraph_text: str) -> None:
    """Helper: create a real, minimal DOCX file with given text."""
    document = docx.Document()
    document.add_paragraph(paragraph_text)
    document.save(path)


def test_read_docx_text_extracts_paragraphs(tmp_path):
    """A valid DOCX's paragraph text should be extracted correctly."""
    docx_path = tmp_path / "sample.docx"
    make_test_docx(docx_path, "Hello DesktopAI from Word")

    result = read_docx_text(docx_path)

    assert result is not None
    assert "Hello DesktopAI from Word" in result


def test_read_docx_text_extracts_table_content(tmp_path):
    """Text inside tables should also be extracted, not just paragraphs."""
    docx_path = tmp_path / "with_table.docx"
    document = docx.Document()
    document.add_paragraph("Intro text")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "CellA"
    table.cell(0, 1).text = "CellB"
    document.save(docx_path)

    result = read_docx_text(docx_path)

    assert result is not None
    assert "CellA" in result
    assert "CellB" in result


def test_read_docx_text_returns_none_for_missing_file(tmp_path):
    """A path that doesn't exist should return None, not raise."""
    missing_path = tmp_path / "does_not_exist.docx"

    result = read_docx_text(missing_path)

    assert result is None


def test_read_docx_text_returns_none_for_corrupt_file(tmp_path):
    """A file that isn't actually a valid DOCX should return None,
    not crash the caller."""
    fake_docx_path = tmp_path / "corrupt.docx"
    fake_docx_path.write_bytes(b"this is not a real docx")

    result = read_docx_text(fake_docx_path)

    assert result is None